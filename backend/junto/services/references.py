from __future__ import annotations

from abc import ABC, abstractmethod
from io import BytesIO
from pathlib import PurePath
from zipfile import ZipFile

from junto.domain.errors import DomainError, invalid


class ReferenceTextExtractor(ABC):
    @abstractmethod
    def extract(self, *, file_name: str, content: bytes) -> tuple[str, str]:
        """Return normalized content type and extracted plain text."""
        raise NotImplementedError


class DefaultReferenceTextExtractor(ReferenceTextExtractor):
    _TEXT_EXTENSIONS = {".txt": "text/plain", ".md": "text/markdown"}

    def __init__(
        self,
        *,
        max_characters: int,
        max_pdf_pages: int = 120,
        max_docx_uncompressed_bytes: int = 20 * 1024 * 1024,
    ) -> None:
        self._max_characters = max_characters
        self._max_pdf_pages = max_pdf_pages
        self._max_docx_uncompressed_bytes = max_docx_uncompressed_bytes

    def extract(self, *, file_name: str, content: bytes) -> tuple[str, str]:
        suffix = PurePath(file_name).suffix.lower()
        if suffix == ".doc":
            raise invalid(
                "UNSUPPORTED_REFERENCE_TYPE",
                "Legacy .doc files are not supported. Save the file as .docx, PDF, or text.",
            )
        if suffix in self._TEXT_EXTENSIONS:
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError as error:
                raise invalid(
                    "REFERENCE_ENCODING_INVALID",
                    "Text reference files must use UTF-8 encoding.",
                ) from error
            content_type = self._TEXT_EXTENSIONS[suffix]
        elif suffix == ".pdf":
            text = self._extract_pdf(content)
            content_type = "application/pdf"
        elif suffix == ".docx":
            text = self._extract_docx(content)
            content_type = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            raise invalid(
                "UNSUPPORTED_REFERENCE_TYPE",
                "Reference material must be a .txt, .md, .pdf, or .docx file.",
            )

        normalized = "\n".join(line.rstrip() for line in text.splitlines()).strip()
        if not normalized:
            raise invalid(
                "REFERENCE_TEXT_EMPTY",
                "No readable text could be extracted from this reference file.",
            )
        if len(normalized) > self._max_characters:
            raise invalid(
                "REFERENCE_TEXT_TOO_LONG",
                f"Extracted reference text exceeds {self._max_characters} characters.",
            )
        return content_type, normalized

    def _extract_pdf(self, content: bytes) -> str:
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(content))
            if len(reader.pages) > self._max_pdf_pages:
                raise invalid(
                    "REFERENCE_PDF_PAGE_LIMIT",
                    f"PDF reference files can contain at most {self._max_pdf_pages} pages.",
                )
            blocks: list[str] = []
            extracted_characters = 0
            for page in reader.pages:
                block = page.extract_text() or ""
                extracted_characters += len(block)
                if extracted_characters > self._max_characters:
                    raise invalid(
                        "REFERENCE_TEXT_TOO_LONG",
                        f"Extracted reference text exceeds {self._max_characters} characters.",
                    )
                blocks.append(block)
            return "\n\n".join(blocks)
        except DomainError:
            raise
        except Exception as error:  # library exceptions vary by malformed PDF
            raise invalid(
                "REFERENCE_EXTRACTION_FAILED",
                "The PDF could not be read. Try exporting it again or upload text.",
            ) from error

    def _extract_docx(self, content: bytes) -> str:
        try:
            with ZipFile(BytesIO(content)) as archive:
                expanded_size = sum(item.file_size for item in archive.infolist())
            if expanded_size > self._max_docx_uncompressed_bytes:
                raise invalid(
                    "REFERENCE_DOCX_EXPANSION_LIMIT",
                    "The DOCX expands beyond the safe reference-material limit.",
                )

            from docx import Document

            document = Document(BytesIO(content))
            blocks: list[str] = []
            extracted_characters = 0

            def append_block(block: str) -> None:
                nonlocal extracted_characters
                extracted_characters += len(block)
                if extracted_characters > self._max_characters:
                    raise invalid(
                        "REFERENCE_TEXT_TOO_LONG",
                        f"Extracted reference text exceeds {self._max_characters} characters.",
                    )
                blocks.append(block)

            for paragraph in document.paragraphs:
                append_block(paragraph.text)
            for table in document.tables:
                for row in table.rows:
                    append_block("\t".join(cell.text for cell in row.cells))
            return "\n".join(blocks)
        except DomainError:
            raise
        except Exception as error:  # library exceptions vary by malformed archive
            raise invalid(
                "REFERENCE_EXTRACTION_FAILED",
                "The DOCX file could not be read. Try exporting it again or upload text.",
            ) from error
